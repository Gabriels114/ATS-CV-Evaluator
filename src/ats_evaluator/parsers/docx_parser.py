from pathlib import Path

from docx import Document

from ..domain.cv import ParseQuality
from ..errors import UnparseableDocumentError
from .base import DocumentParser, ParsedDocument


class DocxParser:
    def can_parse(self, path: Path) -> bool:
        return path.suffix.lower() == ".docx"

    def parse(self, path: Path) -> ParsedDocument:
        try:
            doc = Document(str(path))
        except Exception as exc:
            raise UnparseableDocumentError(f"Failed to open DOCX '{path.name}': {exc}") from exc

        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        tables_detected = len(doc.tables) > 0
        table_texts: list[str] = []
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    table_texts.append(row_text)

        raw_text = "\n".join(paragraphs)
        if table_texts:
            raw_text += "\n\n" + "\n".join(table_texts)

        if not raw_text.strip():
            raise UnparseableDocumentError(
                f"'{path.name}' contains no extractable text."
            )

        quality = ParseQuality(
            page_count=1,  # DOCX doesn't expose page count without rendering
            has_images_only=False,
            tables_detected=tables_detected,
            estimated_char_density=float(len(raw_text)),
        )
        return ParsedDocument(raw_text=raw_text, quality=quality)


_instance: DocumentParser = DocxParser()
